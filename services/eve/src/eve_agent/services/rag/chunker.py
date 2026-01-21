from __future__ import annotations

import io
import os
import tempfile
from typing import List
import json
import csv

from pypdf import PdfReader
from docx import Document as DocxDocument
import pdfplumber

import tiktoken
from langchain_text_splitters import RecursiveCharacterTextSplitter


def extract_pdf_with_pdfplumber(path: str) -> tuple[str, bool]:
    """
    Extract text from PDF using pdfplumber (lightweight alternative to docling).
    This extracts both text and tables in a readable format.
    Automatically falls back to OCR if text extraction is insufficient.
    
    Returns:
        Tuple of (extracted_text, used_ocr)
    """
    output_parts = []
    
    try:
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text from page
                text = page.extract_text()
                if text and text.strip():
                    output_parts.append(f"--- Page {page_num} ---")
                    output_parts.append(text)
                
                # Extract tables if any
                tables = page.extract_tables()
                if tables:
                    for table_idx, table in enumerate(tables, 1):
                        output_parts.append(f"\n[Table {table_idx} on Page {page_num}]")
                        # Convert table to markdown-like format
                        for row in table:
                            if row:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                output_parts.append(row_text)
                        output_parts.append("")  # Empty line after table
        
        pdfplumber_text = "\n".join(output_parts)
        
        try:
            from ..rag.ocr_service import ocr_service
            final_text, used_ocr = ocr_service.extract_with_fallback(path, pdfplumber_text)
            return final_text, used_ocr
        except ImportError:
            # OCR service not available, return pdfplumber text
            return pdfplumber_text, False
    
    except Exception as e:
        # Fallback to pypdf if pdfplumber fails
        print(f"⚠️ pdfplumber failed, falling back to pypdf: {e}")
        reader = PdfReader(path)
        parts = [page.extract_text() or "" for page in reader.pages]
        pypdf_text = "\n".join(parts)
        
        # Try OCR on pypdf text too
        try:
            from ..rag.ocr_service import ocr_service
            final_text, used_ocr = ocr_service.extract_with_fallback(path, pypdf_text)
            return final_text, used_ocr
        except ImportError:
            return pypdf_text, False


def extract_txt_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_json_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


def extract_csv_text(path: str) -> str:
    rows = []
    with open(path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)


def load_document_text(path: str) -> tuple[str, dict]:
    """
    Load and extract text from various document formats.
    
    Returns:
        Tuple of (text, metadata) where metadata contains info like OCR usage
    """
    ext = os.path.splitext(path)[1].lower()
    metadata = {"ocr_used": False}

    if ext == ".pdf":
        text, used_ocr = extract_pdf_with_pdfplumber(path)
        metadata["ocr_used"] = used_ocr
        return text, metadata
    
    if ext == ".docx":
        # Use pdfplumber-style extraction for docx too
        doc = DocxDocument(path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return text, metadata

    if ext == ".txt":
        return extract_txt_text(path), metadata

    if ext == ".json":
        return extract_json_text(path), metadata

    if ext == ".csv":
        return extract_csv_text(path), metadata

    raise ValueError(f"Unsupported file type: {ext}")


def get_token_encoder(model_name: str = "gpt-4o-mini"):
    try:
        return tiktoken.encoding_for_model(model_name)
    except KeyError:
        return tiktoken.get_encoding("cl100k_base")


def chunk_text_token_based(
    text: str,
    chunk_tokens: int = 350,
    overlap_tokens: int = 50,
    model_name: str = "gpt-4o-mini",
) -> List[str]:
    encoding = get_token_encoder(model_name)

    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        encoding_name=encoding.name,
        chunk_size=chunk_tokens,
        chunk_overlap=overlap_tokens,
        separators=[
            "\n\n",  # paragraphs
            "\n",    # headers/lists
            ". ",    # sentence boundary
            " ",     # word boundary
            ""       # fallback to characters
        ],
    )

    return splitter.split_text(text)


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    
    chunk_tokens = max(100, chunk_size // 4)  # Minimum 100 tokens
    overlap_tokens = max(20, overlap // 4)    # Minimum 20 tokens
    
    # token-based chunking
    return chunk_text_token_based(
        text=text,
        chunk_tokens=chunk_tokens,
        overlap_tokens=overlap_tokens,
        model_name="gpt-4o-mini"  # Using GPT-4 tokenizer as default
    )


async def extract_text_from_file(contents: bytes, filename: str) -> tuple[str, dict]:
    """
    Extract text from file bytes (PDF, DOCX, TXT, JSON, CSV).
    Includes OCR support for scanned/handwritten PDFs.
    
    Returns:
        Tuple of (text, metadata) where metadata contains OCR info
    """
    filename = filename.lower() if filename else "upload"
    metadata = {"ocr_used": False}
    
    # Save to temp file for processing
    suffix = os.path.splitext(filename)[1] or ".tmp"
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(contents)
        temp_path = tmp.name
    
    try:
        text, file_metadata = load_document_text(temp_path)
        cleaned = "\n".join(line.rstrip() for line in text.splitlines())
        metadata.update(file_metadata)
        return cleaned, metadata
        
    except ValueError as e:
        # Fallback for unsupported types
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".pdf":
            reader = PdfReader(io.BytesIO(contents))
            parts = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(parts), metadata
            
        elif ext == ".docx":
            doc = DocxDocument(io.BytesIO(contents))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs), metadata
            
        elif ext == ".txt":
            return contents.decode("utf-8"), metadata
            
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass


def ingest_and_chunk(
    path: str,
    chunk_tokens: int = 350,
    overlap_tokens: int = 50,
    model_name: str = "gpt-4o-mini",
) -> dict:
    """
    Ingest document and chunk it.
    
    Returns:
        Dictionary with path, chunks, and metadata (including OCR info)
    """
    raw_text, metadata = load_document_text(path)
    
    cleaned = "\n".join(line.rstrip() for line in raw_text.splitlines())

    chunks = chunk_text_token_based(
        cleaned,
        chunk_tokens=chunk_tokens,
        overlap_tokens=overlap_tokens,
        model_name=model_name,
    )

    return {
        "path": path,
        "num_chunks": len(chunks),
        "chunks": chunks,
        "metadata": metadata
    }
